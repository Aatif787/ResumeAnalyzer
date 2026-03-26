"""
ResumeAI Generator — Generates optimized resume documents (DOCX/PDF).
"""
import io
import re


class ResumeGenerator:
    """Generate a formatted resume document from analysis data."""

    def __init__(self, data, options=None):
        self.data = data
        self.options = options or {}
        self.fmt = self.options.get('format', 'docx')

    def generate(self):
        """Generate and return a file stream."""
        if self.fmt == 'pdf':
            return self._generate_pdf()
        else:
            return self._generate_docx()

    def _generate_docx(self):
        """Generate a DOCX resume."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style defaults
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Name / Header
            name = self.data.get('name', 'Your Name')
            heading = doc.add_heading(name, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info line
            contact_parts = []
            if self.data.get('email'):
                contact_parts.append(self.data['email'])
            if self.data.get('phone'):
                contact_parts.append(self.data['phone'])
            if self.data.get('linkedin'):
                contact_parts.append(self.data['linkedin'])
            if contact_parts:
                cp = doc.add_paragraph(' | '.join(contact_parts))
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Sections
            sections = self.data.get('sections', [])
            for section in sections:
                title = section.get('title', 'Section')
                doc.add_heading(title, level=1)
                items = section.get('items', [])
                for item in items:
                    if isinstance(item, dict):
                        # Role + Company + Dates
                        role_line = item.get('role', '')
                        company = item.get('company', '')
                        dates = item.get('dates', '')
                        if role_line:
                            p = doc.add_paragraph()
                            run = p.add_run(role_line)
                            run.bold = True
                            if company:
                                p.add_run(f' — {company}')
                            if dates:
                                p.add_run(f'  ({dates})')
                        # Bullet points
                        for bullet in item.get('bullets', []):
                            doc.add_paragraph(bullet, style='List Bullet')
                    elif isinstance(item, str):
                        doc.add_paragraph(item)

            # If no structured sections, dump the raw/improved text
            if not sections:
                text = self.data.get('text', self.data.get('improved_text', ''))
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)

            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return stream

        except ImportError:
            # Fallback: return plain text as bytes
            return self._generate_text()

    def _generate_pdf(self):
        """Generate a simple PDF resume (text-based fallback)."""
        # PDF generation would need reportlab or similar
        # For now, generate a text file with .pdf-like content
        return self._generate_text()

    def _generate_text(self):
        """Fallback text generator."""
        text = self.data.get('text', self.data.get('improved_text', 'No content available.'))
        stream = io.BytesIO()
        stream.write(text.encode('utf-8'))
        stream.seek(0)
        return stream
