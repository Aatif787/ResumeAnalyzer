"""
ResumeAI Parser — Extracts text content from PDF, DOCX, and TXT files.
"""
import io
import os


def extract_text(file_obj):
    """
    Extract text from an uploaded Flask file object.
    Supports PDF, DOCX, and TXT formats.
    """
    filename = getattr(file_obj, 'filename', 'unknown.txt').lower()

    if filename.endswith('.pdf'):
        return _extract_pdf(file_obj)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        return _extract_docx(file_obj)
    elif filename.endswith('.txt'):
        raw = file_obj.read()
        if isinstance(raw, bytes):
            return raw.decode('utf-8', errors='ignore')
        return raw
    else:
        raise ValueError(f"Unsupported file format: {os.path.splitext(filename)[1]}")


def _extract_pdf(file_obj):
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_obj.read()))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        result = "\n\n".join(pages_text).strip()
        if not result:
            raise ValueError("PDF appears to be scanned/image-based. Please use a text-based PDF.")
        return result
    except ImportError:
        raise ValueError("PDF parsing library not available. Install PyPDF2.")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _extract_docx(file_obj):
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_obj.read()))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        result = "\n".join(paragraphs).strip()
        if not result:
            raise ValueError("DOCX appears to be empty.")
        return result
    except ImportError:
        raise ValueError("DOCX parsing library not available. Install python-docx.")
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
